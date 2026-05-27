#!/usr/bin/env python3
"""Build a blind JSO-ready submission package from the canonical LaTeX draft."""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SUB = ROOT / "submission" / "jso"
ASSETS = SUB / "assets"
TITLE = "Effective without warrant: Causal-normative networks and the social life of status"
DOCX = SUB / "effective-without-warrant-jso-blind.docx"
RTF = SUB / "effective-without-warrant-jso-blind.rtf"
BLIND_TEX = SUB / "main.blind.tex"
BLIND_MD = SUB / "main.blind.md"
COUNTS = SUB / "submission-checks.md"
CHECK_DATE = date.today().isoformat()
CHICAGO_CSL = Path(
    "/Users/brettreynolds/Library/Application Support/Mendeley Desktop/"
    "citationStyles-1.0/chicago-author-date.csl"
)


FIG_STYLE = r"""
\documentclass[tikz,border=6pt]{standalone}
\usepackage{xcolor}
\usepackage{tikz}
\usetikzlibrary{arrows.meta, positioning}
\definecolor{xGreen}{RGB}{27,158,119}
\definecolor{xOrange}{RGB}{217,95,2}
\definecolor{xPurple}{RGB}{117,112,179}
\tikzset{
  kind/.style={draw, rounded corners, align=center, font=\small, inner sep=3pt, minimum height=9mm},
  ghost/.style={draw, rounded corners, align=center, font=\small, inner sep=3pt, minimum height=9mm,
                draw=gray, text=gray, dash pattern=on 2pt off 1.5pt},
  flagged/.style={draw, rounded corners, align=center, font=\small, inner sep=3pt, minimum height=9mm,
                draw=black, text=black, dash pattern=on 3pt off 2pt},
  causal/.style={-{Stealth}, semithick},
  ntrans/.style={-{Triangle}, semithick, xOrange},
  constit/.style={xPurple, very thick, double, double distance=1pt},
  recog/.style={-{Stealth}, semithick, dashed, xGreen},
}
\newcommand{\edgelegend}{%
  \begin{scope}[shift={(0,-3.4)}]
    \draw[causal] (0,0) -- ++(0.9,0);    \node[right,font=\footnotesize] at (1.0,0) {causal};
    \draw[constit] (0,-0.6) -- ++(0.9,0); \node[right,font=\footnotesize] at (1.0,-0.6) {constitutive};
    \draw[ntrans] (5.2,0) -- ++(0.9,0);   \node[right,font=\footnotesize] at (6.2,0) {normative-transition};
    \draw[recog] (5.2,-0.6) -- ++(0.9,0); \node[right,font=\footnotesize] at (6.2,-0.6) {recognitional};
  \end{scope}%
}
\begin{document}
"""


FIGURES = {
    "fig-promise": r"""
\begin{tikzpicture}
  \node[kind] (u) at (0,0)    {undertaking};
  \node[kind] (o) at (3.7,0)  {duty};
  \node[kind] (k) at (3.7,2)  {claim-right};
  \node[kind] (r) at (7.4,0)  {recognition};
  \node[kind] (d) at (11.0,0) {demand,\\repair};
  \node[kind] (l) at (0,-2)   {reliance};
  \draw[ntrans]  (u) -- (o);
  \draw[recog]   (o) -- (r);
  \draw[causal]  (r) -- (d);
  \draw[constit] (o) -- (k);
  \draw[causal]  (u) -- (l);
  \edgelegend
\end{tikzpicture}
""",
    "fig-conferral": r"""
\begin{tikzpicture}
  \node[kind]    (u) at (0,0)    {signal,\\ascription};
  \node[kind]    (r) at (3.5,0)  {recognition,\\classification};
  \node[kind]    (p) at (7.1,0)  {standing agents,\\uptake, records};
  \node[flagged] (s) at (11.1,0) {ascribed\\standing};
  \node[kind]    (c) at (11.1,-2) {constraints,\\enablements};
  \node[kind]    (h) at (14.5,-2) {harm};
  \node[font=\small, align=left, text=gray, anchor=west] at (8.8,1.55)
    {moral warrant: absent\\[-1pt]\scriptsize assessment condition, not a node};
  \draw[causal] (u) -- (r);
  \draw[causal] (r) -- (p);
  \draw[ntrans] (p) -- (s);
  \draw[constit] (s) -- (c);
  \draw[causal] (c) -- (h);
  \draw[recog]  (s) to[bend left=18] (r);
  \node[font=\footnotesize\itshape, xGreen, anchor=south] at (9.1,0.7) {effective, conferred};
\end{tikzpicture}
""",
}


CONTRACTIONS = [
    ("doesn't", "does not"),
    ("don't", "do not"),
    ("can't", "cannot"),
    ("isn't", "is not"),
    ("aren't", "are not"),
    ("hasn't", "has not"),
    ("haven't", "have not"),
    ("wasn't", "was not"),
    ("weren't", "were not"),
    ("didn't", "did not"),
    ("won't", "will not"),
    ("wouldn't", "would not"),
    ("shouldn't", "should not"),
    ("couldn't", "could not"),
    ("needn't", "need not"),
    ("what's", "what is"),
    ("that's", "that is"),
    ("there's", "there is"),
]


def run(cmd: list[str], cwd: Path = ROOT, capture: bool = False) -> subprocess.CompletedProcess[str]:
    result = subprocess.run(
        cmd,
        cwd=cwd,
        text=True,
        stdout=subprocess.PIPE if capture else None,
        stderr=subprocess.PIPE if capture else None,
        check=False,
    )
    if result.returncode != 0:
        if capture:
            sys.stderr.write(result.stdout or "")
            sys.stderr.write(result.stderr or "")
        raise SystemExit(f"Command failed ({result.returncode}): {' '.join(cmd)}")
    return result


def set_run_font(run, name: str, size: int) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)


def add_page_number_field(run) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def decontract(text: str) -> str:
    for old, new in CONTRACTIONS:
        text = re.sub(rf"\b{re.escape(old)}\b", new, text, flags=re.IGNORECASE)
    return text


def cleanup_crossrefs(text: str) -> str:
    text = re.sub(r'<a href="[^"]*"[^>]*>(.*?)</a>', r"\1", text)
    text = re.sub(r'\[([0-9]+)\]\([^)]*\)\{reference-type="[^"]+" reference="[^"]+"\}', r"\1", text)
    text = re.sub(r'\[([0-9]+)\]\([^)]*\)\{[^}]*\}', r"\1", text)
    text = re.sub(r'\s*\{reference-type="[^"]+" reference="[^"]+"\}', "", text)
    return text


def strip_html(text: str) -> str:
    text = cleanup_crossrefs(text)
    text = re.sub(r"<[^>]+>", "", text)
    return re.sub(r"\s+", " ", text).strip()


def make_blind_tex() -> None:
    tex = (ROOT / "main.tex").read_text(encoding="utf-8")
    tex = re.sub(
        r"(\\hypersetup\{\s*pdftitle=\{[^{}]*\})(\s*\})",
        r"\1,\n  pdfauthor={}\2",
        tex,
        flags=re.S,
    )
    tex = re.sub(r"\\author\{Brett Reynolds.*?\\date\{\\today\}", r"\\author{}\n\\date{}", tex, flags=re.S)
    tex = re.sub(
        r"\n\\section\*\{Acknowledgements\}.*?(?=\n\\newpage\s*\n\\printbibliography)",
        "\n",
        tex,
        flags=re.S,
    )
    tex = re.sub(r"\n% TODO:.*", "", tex)
    BLIND_TEX.write_text(tex, encoding="utf-8")


def build_figures() -> None:
    for stem, body in FIGURES.items():
        tex_path = ASSETS / f"{stem}.tex"
        tex_path.write_text(f"{FIG_STYLE}\n{body}\n\\end{{document}}\n", encoding="utf-8")
        run(["xelatex", "-interaction=nonstopmode", "-halt-on-error", tex_path.name], cwd=ASSETS, capture=True)
        run(["pdfcrop", f"{stem}.pdf", f"{stem}-crop.pdf"], cwd=ASSETS, capture=True)
        run(["pdftoppm", "-png", "-singlefile", "-r", "300", f"{stem}-crop.pdf", stem], cwd=ASSETS, capture=True)


def abstract_and_keywords() -> tuple[str, str]:
    raw = (ROOT / "abstract.md").read_text(encoding="utf-8")
    body = raw.split("\n\n", 1)[1].strip()
    abstract, keywords = body.split("\n\nKeywords:", 1)
    return decontract(abstract.strip()), keywords.strip()


def make_markdown() -> None:
    result = run(["pandoc", "--from=latex", "--to=markdown", "--wrap=none", str(BLIND_TEX)], capture=True)
    md = result.stdout
    md = re.sub(r'^\s*(?:\*\*Keywords:\*\*|Keywords:).*?\n+(?=# )', "", md, count=1, flags=re.S)
    md = re.sub(r"\n# Acknowledgements.*?(?=\n\[\^|\Z)", "\n", md, flags=re.S)
    md = re.sub(r"\n\[\^1\]: Contact:.*", "", md)

    def replace_figure(match: re.Match[str], image: str, number: int) -> str:
        block = match.group(0)
        caption_match = re.search(r"<figcaption>(.*?)</figcaption>", block, flags=re.S)
        caption = strip_html(caption_match.group(1) if caption_match else "")
        return f"\n\n![Figure {number}. {caption}](assets/{image}.png){{width=6.5in}}\n\n"

    md = re.sub(
        r'<figure id="fig:promise".*?</figure>',
        lambda m: replace_figure(m, "fig-promise", 1),
        md,
        flags=re.S,
    )
    md = re.sub(
        r'<figure id="fig:misrecognition".*?</figure>',
        lambda m: replace_figure(m, "fig-conferral", 2),
        md,
        flags=re.S,
    )
    md = cleanup_crossrefs(md)
    md = decontract(md)

    abstract, keywords = abstract_and_keywords()
    front = f"""---
reference-section-title: "References"
---

# {TITLE} {{.unnumbered}}

# Abstract {{.unnumbered}}

{abstract}

Keywords: {keywords}

"""
    BLIND_MD.write_text(front + md, encoding="utf-8")


def build_pdf() -> None:
    run(
        [
            "xelatex",
            "-interaction=nonstopmode",
            "-halt-on-error",
            "-jobname=jso-blind",
            "-output-directory",
            str(SUB),
            str(BLIND_TEX),
        ],
        capture=True,
    )
    run(["biber", str(SUB / "jso-blind")], capture=True)
    for _ in range(2):
        run(
            [
                "xelatex",
                "-interaction=nonstopmode",
                "-halt-on-error",
                "-jobname=jso-blind",
                "-output-directory",
                str(SUB),
                str(BLIND_TEX),
            ],
            capture=True,
        )


def pandoc_args(output: Path) -> list[str]:
    args = [
        "pandoc",
        str(BLIND_MD.name),
        "--from=markdown+citations",
        "--citeproc",
        "--bibliography",
        str((ROOT / "references-standalone.bib").resolve()),
        "-o",
        str(output.name),
    ]
    if CHICAGO_CSL.exists():
        args.extend(["--csl", str(CHICAGO_CSL)])
    return args


def patch_docx() -> None:
    doc = Document(DOCX)
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.title = TITLE
    props.subject = ""
    props.keywords = ""
    props.comments = ""

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        footer_para = section.footer.paragraphs[0]
        footer_para._p.clear_content()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run()
        set_run_font(run, "Times New Roman", 12)
        add_page_number_field(run)

    for paragraph in list(doc.paragraphs):
        if paragraph.style and paragraph.style.name == "Author" and not paragraph.text.strip():
            paragraph._element.getparent().remove(paragraph._element)

    for style_name in ["Normal", "Body Text", "Default Paragraph Font"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            if hasattr(style, "font"):
                style.font.name = "Times New Roman"
                style.font.size = Pt(12)
            if hasattr(style, "paragraph_format"):
                style.paragraph_format.line_spacing = 1
                style.paragraph_format.space_after = Pt(6)

    for heading in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if heading in doc.styles:
            style = doc.styles[heading]
            style.font.name = "Times New Roman"
            style.font.size = Pt(14 if heading == "Title" else 12)
            style.font.bold = heading in {"Title", "Heading 1"}
            style.font.color.rgb = None

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1
        for run in paragraph.runs:
            set_run_font(run, "Times New Roman", 12)

    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1
                    paragraph.paragraph_format.space_after = Pt(3)
                    for run in paragraph.runs:
                        set_run_font(run, "Times New Roman", 10)

    doc.save(DOCX)


def scrub_docx_package(path: Path) -> None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in {"docProps/custom.xml", "word/comments.xml"}:
                continue
            data = zin.read(item.filename)
            if item.filename == "_rels/.rels":
                data = re.sub(
                    rb'<Relationship[^>]+Type="[^"]*/custom-properties"[^>]*/>',
                    b"",
                    data,
                )
            elif item.filename == "word/_rels/document.xml.rels":
                data = re.sub(
                    rb'<Relationship[^>]+Type="[^"]*/comments"[^>]*/>',
                    b"",
                    data,
                )
            elif item.filename == "[Content_Types].xml":
                data = re.sub(
                    rb'<Override[^>]+PartName="/docProps/custom.xml"[^>]*/>',
                    b"",
                    data,
                )
                data = re.sub(
                    rb'<Override[^>]+PartName="/word/comments.xml"[^>]*/>',
                    b"",
                    data,
                )
            zout.writestr(item, data)
    shutil.move(tmp_path, path)
    path.chmod(0o644)


def build_word_outputs() -> None:
    run(pandoc_args(DOCX), cwd=SUB, capture=True)
    patch_docx()
    scrub_docx_package(DOCX)
    run(pandoc_args(RTF), cwd=SUB, capture=True)


def plain_text(path: Path) -> str:
    return run(["pandoc", str(path), "--to=plain", "--wrap=none"], capture=True).stdout


def write_checks() -> None:
    plain = plain_text(BLIND_MD)
    plain = re.sub(r"^Effective without warrant:.*?\n", "", plain, flags=re.S)
    body = plain.split("References", 1)[0]
    body_no_abstract = body.split("The problem of mixed explanation", 1)[-1]
    char_count = len(body_no_abstract)
    word_count = len(re.findall(r"\b\w+\b", body_no_abstract))
    abstract = abstract_and_keywords()[0]
    abstract_words = len(re.findall(r"\b\w+\b", abstract))
    text = f"""# JSO Submission Checks

- Working venue: Journal of Social Ontology regular article.
- Submission requirements checked against: https://journalofsocialontology.org/index.php/jso/submission ({CHECK_DATE}).
- Working title: {TITLE}
- Blind manuscript DOCX: `{DOCX.name}`
- Blind manuscript RTF: `{RTF.name}`
- Blind PDF for internal checking: `jso-blind.pdf`
- Abstract words: {abstract_words} (JSO limit: 250)
- Body/footnote characters, approximate from blind Markdown excluding abstract and references: {char_count} (JSO regular-article limit: 60,000 including spaces)
- Body words, approximate from blind Markdown excluding abstract and references: {word_count}
- Font/layout target for DOCX: Times New Roman 12 pt, single-spaced body, 1 inch margins; tables use 10 pt to avoid clipping.
- Citation target: Chicago author-date via Pandoc citeproc and local `chicago-author-date.csl`.
- Anonymization: author line, affiliation, contact footnote, PDF author metadata, and acknowledgements removed in blind source.
"""
    COUNTS.write_text(text, encoding="utf-8")


def package_zip() -> None:
    for old_zip in SUB.glob("effective-without-warrant-jso-blind-package-*.zip"):
        old_zip.unlink()
    zip_path = SUB / f"effective-without-warrant-jso-blind-package-{CHECK_DATE}.zip"
    members = [DOCX, RTF, SUB / "jso-blind.pdf", COUNTS]
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for member in members:
            zf.write(member, member.name)


def cleanup_intermediates() -> None:
    for pattern in ["jso-blind.aux", "jso-blind.bbl", "jso-blind.bcf", "jso-blind.blg", "jso-blind.log", "jso-blind.out", "jso-blind.run.xml"]:
        path = SUB / pattern
        if path.exists():
            path.unlink()
    for path in ASSETS.glob("fig-*"):
        if path.suffix in {".aux", ".log", ".pdf", ".tex"} or path.name.endswith("-crop.pdf"):
            path.unlink()


def main() -> None:
    SUB.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    make_blind_tex()
    build_figures()
    make_markdown()
    build_pdf()
    build_word_outputs()
    write_checks()
    package_zip()
    cleanup_intermediates()
    print(f"Built {DOCX.relative_to(ROOT)}")
    print(f"Built {RTF.relative_to(ROOT)}")
    print(f"Built {(SUB / 'jso-blind.pdf').relative_to(ROOT)}")
    print(f"Built {COUNTS.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
